# 实验四 Markdown 转 Word 脚本
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 配置路径
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(REPORT_DIR, '实验报告草稿.md')
OUTPUT_FILE = os.path.join(REPORT_DIR, '实验四报告.docx')
PICTURES_DIR = os.path.join(REPORT_DIR, 'pictures')

def set_chinese_font(run, font_name='SimSun', font_size=Pt(12), bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold

def add_heading_chinese(doc, text, level=1):
    """添加中文标题"""
    if level == 0:
        # 文档标题
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_chinese_font(run, 'SimHei', Pt(18), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        heading = doc.add_heading(level=level)
        run = heading.runs[0] if heading.runs else heading.add_run()
        run.text = text
        font_name = 'SimHei' if level <= 2 else 'SimSun'
        set_chinese_font(run, font_name, Pt(14 if level == 1 else 12), bold=True)
    return p if level == 0 else heading

def add_paragraph_chinese(doc, text, bold=False, indent=True):
    """添加中文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', Pt(12), bold)
    return p

def process_table(doc, lines, start_idx):
    """处理 Markdown 表格"""
    table_lines = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return idx
    
    # 解析表头
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    # 跳过分隔行
    data_rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            data_rows.append(cells)
    
    if not headers:
        return idx
    
    # 创建表格
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimHei', Pt(11), bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置数据行
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_text
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, 'SimSun', Pt(11))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置列宽
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(8)
    
    return idx

def insert_image(doc, image_path, caption=None):
    """插入图片"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            set_chinese_font(run, 'SimSun', Pt(10), bold=False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print(f"警告: 图片不存在: {image_path}")

def parse_markdown(md_file):
    """解析 Markdown 文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # 默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    i = 0
    code_block = False
    code_content = []
    
    # 图片映射（根据报告内容插入对应的图片）
    # 图1: Ollama 服务截图 - 使用第二张截图（16-18-57，较大）
    # 图2: OpenCode 启动截图 - 使用包含 opencode 的截图
    image_map = {
        '图1': os.path.join(PICTURES_DIR, 'PixPin_2026-06-10_16-45-04.png'),  # Ollama 服务截图
        '图2': os.path.join(PICTURES_DIR, 'PixPin_2026-06-10_16-18-57.png'),  # OpenCode 启动截图
    }
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if code_block:
                # 结束代码块
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                set_chinese_font(run, 'Courier New', Pt(10))
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.line_spacing = 1.15
                code_block = False
                code_content = []
            else:
                code_block = True
            i += 1
            continue
        
        if code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            text = line[2:]
            add_heading_chinese(doc, text, level=0)
            i += 1
            continue
        elif line.startswith('## '):
            text = line[3:]
            add_heading_chinese(doc, text, level=1)
            i += 1
            continue
        elif line.startswith('### '):
            text = line[4:]
            add_heading_chinese(doc, text, level=2)
            i += 1
            continue
        
        # 处理表格
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            i = process_table(doc, lines, i)
            continue
        elif line.strip().startswith('|'):
            # 单个表格行（无表头的情况）
            i += 1
            continue
        
        # 处理引用/提示块
        if line.startswith('> '):
            text = line[2:]
            # 检查是否是图片标记
            if '图1' in text and 'Ollama' in text:
                insert_image(doc, image_map['图1'], '图1 Ollama 服务检测')
                i += 1
                continue
            elif '图2' in text and 'OpenCode' in text:
                insert_image(doc, image_map['图2'], '图2 OpenCode 启动界面')
                i += 1
                continue
            
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, 'SimSun', Pt(11), bold=False)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue
        
        # 处理列表
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # 检查是否包含复选框
            if text.startswith('✅ '):
                text = text[2:]
            
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            set_chinese_font(run, 'SimSun', Pt(12))
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue
        
        # 处理普通段落（删除粗体/斜体标记，统一普通文本）
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)
        
        # 删除 ** 和 * 标记，直接作为普通文本
        text = line.replace('**', '').replace('*', '')
        run = p.add_run(text)
        set_chinese_font(run, 'SimSun', Pt(12))
        
        i += 1
    
    return doc

def main():
    print("开始生成 Word 文档...")
    
    if not os.path.exists(MD_FILE):
        print(f"错误: 找不到 Markdown 文件: {MD_FILE}")
        return
    
    # 解析 Markdown 并生成 Word
    doc = parse_markdown(MD_FILE)
    
    # 保存文档
    doc.save(OUTPUT_FILE)
    print(f"Word 文档已生成: {OUTPUT_FILE}")
    print(f"文件大小: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")

if __name__ == '__main__':
    main()
